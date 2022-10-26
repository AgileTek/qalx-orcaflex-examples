from io import BytesIO

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pyqalx import Bot
from pyqalx.bot import QalxJob

# Create bot with name `RiserReportBot`
riser_bot = Bot("RiserReportBot")


def write_batch_report(batch, session):
    """
    Take a batch and write a report containing:
        - a title based on the batch title
        - a table of the load cases
        - a summary of the extracted range graphs

    The document will be saved back to the set and a notification will be sent to the
    person who created the batch with a link to the report.

    :param batch: QalxOrcaFlex batch
    :param session: qalx session
    :return: Riser report python-docx document
    """
    report = Document()
    if not batch.meta.get("results_summary"):
        report_title = f"{batch.meta['name']} does not have a results summary."
        report.add_heading(report_title, level=0)
    else:
        report_title = f"Riser report for batch: {batch.meta['name']}"
        report.add_heading(report_title, level=0)
        report.add_paragraph(
            f"This report contains a summary of analysis {report_title}."
        )

        # LOAD CASE DATA
        report.add_heading("Load case data", level=2)
        report.add_paragraph("The load cases examined are detailed in table 1 below.")
        lc_table = report.add_table(rows=len(batch.sets) + 1, cols=3)
        lc_table.style = "LightShading-Accent1"
        headers = lc_table.rows[0].cells
        headers[0].text = "Case Number"
        headers[1].text = "Offset"
        headers[2].text = "Direction"
        for n, (_, case_guid) in enumerate(batch.sets.items()):
            case = session.set.get(case_guid, fields=["items.load_case_info"])
            load_case_info = case["items"].load_case_info.data.raw_info
            row = lc_table.rows[n + 1].cells
            row[0].text = str(n + 1)
            row[1].text = str(load_case_info[0]["value"])
            row[2].text = str(load_case_info[1]["value"])
        cap = report.add_paragraph("Table 1: Load cases")
        cap.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # RESULTS SUMMARY
        results_summary = session.item.get(batch.meta.get("results_summary"))
        report.add_heading("Results summary", level=2)
        report.add_paragraph(
            f"There were {len(results_summary.data['Range Graphs'])} results processed."
        )
        res_n = 1
        for _, result in results_summary.data["Range Graphs"].items():
            report.add_heading(result.result_meta.full_name, level=3)
            res_table = report.add_table(rows=5, cols=5)
            res_table.style = "LightShading-Accent1"
            headers = res_table.rows[0].cells
            headers[0].text = "Stat"
            headers[1].text = "Value"
            headers[2].text = "Arc"
            headers[3].text = "Offset"
            headers[4].text = "Direction"
            for r, stat in enumerate(["max", "min", "static_max", "static_min"]):
                row = res_table.rows[r + 1].cells
                row[0].text = stat.upper()
                value = result[f"{stat}_value"]
                row[1].text = f"{value:2.2f} {result.result_meta.units}"
                arc = result[f"arc_{stat}_value"]
                row[2].text = f"{arc:2.2f}m"
                lci = result[f"{stat}_case_info"]
                row[3].text = lci["offset"]
                row[4].text = lci["direction"]
            cap = report.add_paragraph(
                f"Table {res_n + 1}: {result.result_meta.full_name}"
            )
            cap.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            res_n += 1
        return report


# `process` decorator to identify this method as the relevant step function
# https://docs.qalx.net/bots#process
@riser_bot.process
def write_riser(job: QalxJob):
    batch = job.entity
    report = write_batch_report(batch, job.session)
    report_stream = BytesIO()
    report.save(report_stream)
    report_stream.seek(0)
    batch.meta["report"] = job.session.item.add(
        source=report_stream, file_name=f"{batch.meta['name']}.docx"
    )
    job.save_entity()
    msg_html = f"""
    <html><h1>{batch.meta['name']}</h1></html>

    <p>Report is available to <a href="{batch.meta['report']['file']['url']}">
    download</a>

    """
    job.session.notification.add(
        subject=batch.meta["name"],
        to=[job.entity.info.created.by.email],
        message=msg_html,
    )
